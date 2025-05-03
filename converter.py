import sys
import os
import re
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def read_markdown_file(file_path):
    """Read content from a markdown file."""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def copy_template(template_path, output_path):
    """Copy the template DOCX file to a new file."""
    shutil.copy(template_path, output_path)

def get_style_name(doc, default_name, style_type=WD_STYLE_TYPE.PARAGRAPH):
    """Try to find an appropriate style name in the document."""
    # Special handling for list styles which are often stored as paragraph styles in templates
    if default_name in ['List Bullet', 'List Number'] and style_type == WD_STYLE_TYPE.PARAGRAPH:
        # Just return the name, we'll handle it specially when applying
        return default_name
    
    # Check if the exact style exists
    try:
        doc.styles.get_style_id(default_name, style_type)
        return default_name
    except KeyError:
        pass
    except ValueError:
        # Style exists but wrong type, just return None
        return None
    
    # Common variations for English and other languages
    variations = {
        'Heading1': ['Heading 1', '标题 1', '标题1', 'Heading1', '1', 'Title'],
        'Heading2': ['Heading 2', '标题 2', '标题2', 'Heading2', '2'],
        'Heading3': ['Heading 3', '标题 3', '标题3', 'Heading3', '3'],
        'Heading4': ['Heading 4', '标题 4', '标题4', 'Heading4', '4'],
        'Caption': ['Caption', '标题', '图表标题', 'Table Caption', 'Figure Caption'],
        'List Bullet': ['List Bullet', '项目符号', 'Bulleted List', '无序列表'],
        'List Number': ['List Number', '编号', 'Numbered List', '有序列表'],
        'Normal': ['Normal', '正文', 'Body Text', 'Regular'],
        'No Spacing': ['No Spacing', '无间距', 'Compact'],
        'Intense Quote': ['Intense Quote', '强调引用', 'Strong Quote'],
        'Quote': ['Quote', '引用', 'Block Quote']
    }
    
    # Try to find a matching style
    if default_name in variations:
        for style_name in variations[default_name]:
            try:
                doc.styles.get_style_id(style_name, style_type)
                return style_name
            except (KeyError, ValueError):
                continue
    
    # Fallback to Normal style for paragraphs
    if style_type == WD_STYLE_TYPE.PARAGRAPH:
        try:
            doc.styles.get_style_id('Normal', style_type)
            return 'Normal'
        except (KeyError, ValueError):
            pass
    
    # If nothing works, return None and handle it later
    return None

def find_table_style(doc):
    """Find a suitable table style."""
    table_styles = ['Table Grid', 'Grid Table 1 Light', 'Plain Table 1', 'Table Normal']
    for style_name in table_styles:
        try:
            # Just check if the style exists
            if style_name in doc.styles:
                return style_name
        except:
            continue
    return None

def add_list_item(doc, text, is_bullet=True, number=None):
    """Add a list item with proper formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    
    if is_bullet:
        run = p.add_run("• " + text)
    else:
        run = p.add_run(f"({number}) {text}" if number else f"{text}")
    
    return p

def convert_markdown_to_docx(markdown_content, template_path, output_path):
    """Convert markdown content to DOCX using the template as a base."""
    # Copy template to output file
    copy_template(template_path, output_path)
    
    # Create a new document from the template
    doc = Document(output_path)
    
    # Save original style definitions
    styles = {}
    for style_type in ['Heading1', 'Heading2', 'Heading3', 'Heading4', 'Caption', 
                      'Normal', 'No Spacing', 'Intense Quote', 'Quote']:
        styles[style_type] = get_style_name(doc, style_type)
    
    # Find a suitable table style
    table_style = find_table_style(doc)
    
    # Clear existing content but preserve styles
    paragraphs_to_remove = []
    for i, paragraph in enumerate(doc.paragraphs):
        paragraphs_to_remove.append(i)
    
    # Remove paragraphs in reverse order to avoid index issues
    for i in reversed(paragraphs_to_remove):
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
    
    # Remove all tables
    for table in doc.tables:
        element = table._element
        element.getparent().remove(element)
    
    # Split content into lines
    lines = markdown_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Process headings
        if line.startswith('# '):  # Chapter title
            p = doc.add_paragraph(line[2:])
            if styles['Heading1']:
                try:
                    p.style = styles['Heading1']
                except:
                    run = p.runs[0] if p.runs else p.add_run()
                    run.bold = True
                    run.font.size = Pt(16)
            else:
                run = p.runs[0] if p.runs else p.add_run()
                run.bold = True
                run.font.size = Pt(16)
            i += 1
            continue
        
        if line.startswith('## '):  # Section title
            p = doc.add_paragraph(line[3:])
            if styles['Heading2']:
                try:
                    p.style = styles['Heading2']
                except:
                    run = p.runs[0] if p.runs else p.add_run()
                    run.bold = True
                    run.font.size = Pt(14)
            else:
                run = p.runs[0] if p.runs else p.add_run()
                run.bold = True
                run.font.size = Pt(14)
            i += 1
            continue
        
        if line.startswith('### '):  # Subsection title
            p = doc.add_paragraph(line[4:])
            if styles['Heading3']:
                try:
                    p.style = styles['Heading3']
                except:
                    run = p.runs[0] if p.runs else p.add_run()
                    run.bold = True
                    run.font.size = Pt(13)
            else:
                run = p.runs[0] if p.runs else p.add_run()
                run.bold = True
                run.font.size = Pt(13)
            i += 1
            continue
        
        if line.startswith('#### '):  # Fourth level title
            p = doc.add_paragraph(line[5:])
            if styles['Heading4']:
                try:
                    p.style = styles['Heading4']
                except:
                    run = p.runs[0] if p.runs else p.add_run()
                    run.bold = True
                    run.font.size = Pt(12)
            else:
                run = p.runs[0] if p.runs else p.add_run()
                run.bold = True
                run.font.size = Pt(12)
            i += 1
            continue
        
        # Process code blocks
        if line.startswith('```'):
            # Start of code block
            code_lang = line[3:].strip()  # Get language if specified
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            
            # Check for code listing title
            if i < len(lines) and re.match(r'^代码清单\d+-\d+', lines[i]):
                p = doc.add_paragraph(lines[i])
                if styles['Caption']:
                    try:
                        p.style = styles['Caption']
                    except:
                        run = p.runs[0] if p.runs else p.add_run()
                        run.italic = True
                else:
                    run = p.runs[0] if p.runs else p.add_run()
                    run.italic = True
                i += 1
            
            # Add code block
            for code_line in code_lines:
                p = doc.add_paragraph(code_line)
                if styles['No Spacing']:
                    try:
                        p.style = styles['No Spacing']
                    except:
                        pass
                
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            
            continue
        
        # Process aside sections (【避坑指南】)
        if line.strip() == "<aside>" or "【避坑指南】" in line:
            if line.strip() == "<aside>":
                i += 1  # Move to next line that should contain 【避坑指南】
                if i < len(lines) and "【避坑指南】" in lines[i]:
                    line = lines[i]
                else:
                    continue  # Skip malformed aside
            
            # Create a box for the 【避坑指南】 section
            table = doc.add_table(rows=1, cols=1)
            if table_style:
                try:
                    table.style = table_style
                except:
                    pass
            
            # Style the table as a box
            cell = table.cell(0, 0)
            
            # Process the 【避坑指南】 line
            p = cell.paragraphs[0]
            run = p.add_run(line)
            run.bold = True
            i += 1
            
            # Process content inside the aside
            while i < len(lines):
                if lines[i].strip() == "</aside>" or lines[i].strip().startswith("#") or "【避坑指南】" in lines[i]:
                    break
                
                if lines[i].strip():
                    p = cell.add_paragraph(lines[i])
                i += 1
            
            # Skip </aside> tag if present
            if i < len(lines) and lines[i].strip() == "</aside>":
                i += 1
            
            continue
        
        # Process bullet lists
        if line.strip().startswith('- '):
            items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                items.append(lines[i].strip()[2:])
                i += 1
            
            for item in items:
                # Always use manual list formatting to avoid style type issues
                add_list_item(doc, item, is_bullet=True)
            
            continue
        
        # Process numbered lists with parentheses (X)
        if re.match(r'^\(\d+\)', line.strip()):
            items = []
            numbers = []
            while i < len(lines) and re.match(r'^\(\d+\)', lines[i].strip()):
                item_match = re.match(r'^\((\d+)\)\s*(.+)$', lines[i].strip())
                if item_match:
                    numbers.append(item_match.group(1))
                    items.append(item_match.group(2))
                i += 1
            
            for idx, item in enumerate(items):
                # Always use manual list formatting
                add_list_item(doc, item, is_bullet=False, number=numbers[idx])
            
            continue
        
        # Process tables
        if line.strip().startswith('|') and '|' in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                if not lines[i].strip().startswith('|-'):  # Skip separator lines
                    table_lines.append(lines[i])
                i += 1
                # Skip separator line
                if i < len(lines) and lines[i].strip().startswith('|-'):
                    i += 1
            
            # Skip empty lines after table
            while i < len(lines) and not lines[i].strip():
                i += 1
            
            # Check for table caption
            if i < len(lines) and lines[i].startswith('>') and '表' in lines[i]:
                p = doc.add_paragraph(lines[i].strip('> '))
                if styles['Caption']:
                    try:
                        p.style = styles['Caption']
                    except:
                        p.italic = True
                else:
                    p.italic = True
                i += 1
            
            # Process table content
            if len(table_lines) >= 1:  # Allow tables with just a header
                header_row = table_lines[0]
                data_rows = table_lines[1:] if len(table_lines) > 1 else []
                
                # Parse header row
                header_cells = [cell.strip() for cell in header_row.split('|')[1:-1] if cell.strip()]
                
                if not header_cells:  # Handle malformed table
                    continue
                
                # Create table
                table = doc.add_table(rows=1, cols=len(header_cells))
                if table_style:
                    try:
                        table.style = table_style
                    except:
                        pass
                
                # Add header
                for j, cell_text in enumerate(header_cells):
                    table.cell(0, j).text = cell_text
                
                # Add data rows
                for row_text in data_rows:
                    cells = [cell.strip() for cell in row_text.split('|')[1:-1]]
                    if not cells:
                        continue
                    
                    row_cells = table.add_row().cells
                    for j, cell_text in enumerate(cells[:len(header_cells)]):
                        row_cells[j].text = cell_text
            
            continue
        
        # Process image references
        if line.startswith('!['):
            img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if img_match:
                # Create a placeholder for the image
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.add_run("[Image placeholder: " + img_match.group(2) + "]")
            i += 1
            
            # Check for image caption
            if i < len(lines) and re.match(r'^图\d+-\d+', lines[i]):
                p = doc.add_paragraph(lines[i])
                if styles['Caption']:
                    try:
                        p.style = styles['Caption']
                    except:
                        p.italic = True
                else:
                    p.italic = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                i += 1
            
            continue
        
        # Process notes (注意, 提示)
        if line.strip() in ["注意", "提示"]:
            p = doc.add_paragraph(line.strip())
            if styles['Intense Quote']:
                try:
                    p.style = styles['Intense Quote']
                except:
                    run = p.runs[0] if p.runs else p.add_run()
                    run.bold = True
            else:
                run = p.runs[0] if p.runs else p.add_run()
                run.bold = True
            i += 1
            
            note_content = []
            while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('<'):
                note_content.append(lines[i])
                i += 1
            
            # Create a box for note content
            if note_content:
                table = doc.add_table(rows=1, cols=1)
                if table_style:
                    try:
                        table.style = table_style
                    except:
                        pass
                
                cell = table.cell(0, 0)
                
                for note_line in note_content:
                    if not cell.paragraphs[0].text:
                        p = cell.paragraphs[0]
                    else:
                        p = cell.add_paragraph()
                    
                    p.add_run(note_line)
                    if styles['Quote']:
                        try:
                            p.style = styles['Quote']
                        except:
                            pass
            
            continue
        
        # Process regular paragraphs
        if line.strip():
            p = doc.add_paragraph(line)
            if styles['Normal']:
                try:
                    p.style = styles['Normal']
                except:
                    pass
            i += 1
            continue
        
        # Skip empty lines
        i += 1
    
    # Save the document
    doc.save(output_path)
    print(f"Conversion complete! Document saved to {output_path}")

def main():
    if len(sys.argv) != 2:
        print("Usage: python converter.py <markdown_file>")
        sys.exit(1)
    
    markdown_file = sys.argv[1]
    template_path = "ch04-to-template.docx"
    output_path = "ch04-to.docx"
    
    # Check if files exist
    if not os.path.exists(markdown_file):
        print(f"Error: {markdown_file} does not exist")
        sys.exit(1)
    
    if not os.path.exists(template_path):
        print(f"Error: {template_path} does not exist")
        sys.exit(1)
    
    # Read markdown content
    content = read_markdown_file(markdown_file)
    
    # Convert markdown to DOCX
    convert_markdown_to_docx(content, template_path, output_path)
    
    print(f"\nNote: The converter has made its best attempt to match the formatting.")
    print(f"      Please open '{output_path}' to verify the result.")
    print(f"      Special sections like '【避坑指南】' have been preserved with appropriate styling.")

if __name__ == "__main__":
    main()